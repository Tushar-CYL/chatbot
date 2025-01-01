# import streamlit as st
# import google.generativeai as genai
# from gtts import gTTS
# from pydub import AudioSegment
# from pydub.playback import play
# import tempfile
# import os
# import fitz  # PyMuPDF for PDF files
# import docx  # python-docx for DOCX files
# from datetime import datetime

# GOOGLE_API_KEY = 'AIzaSyDLlgyR6PMeaJPRvLMBqezBCa9HIvvfu8Q'
# genai.configure(api_key=GOOGLE_API_KEY)

# st.set_page_config(page_title="Chatlop", page_icon="🤖", layout="wide")

# prompt = [
# """
# You are an expert in providing detailed and accurate answers to questions based on your vast knowledge. Answer the following questions to the best of your ability.
# """
# ]

# def get_gemini_response(question, prompt):
#     try:
#         model = genai.GenerativeModel('gemini-pro')
#         response = model.generate_content([prompt[0], question])
#         return response.text
#     except Exception as e:
#         st.error(f"Error getting response from GenAI: {e}")
#         return None

# def read_uploaded_file(file):
#     try:
#         if file is not None:
#             file_type = file.type
#             if 'pdf' in file_type:
#                 return read_pdf(file)
#             elif 'wordprocessingml' in file_type:
#                 return read_docx(file)
#             else:
#                 st.error("Unsupported file type.")
#                 return None
#         return None
#     except Exception as e:
#         st.error(f"Error reading uploaded file: {e}")
#         return None

# def read_pdf(file):
#     try:
#         pdf_document = fitz.open(stream=file.read(), filetype="pdf")
#         text = ""
#         for page_num in range(len(pdf_document)):
#             page = pdf_document.load_page(page_num)
#             text += page.get_text()
#         return text
#     except Exception as e:
#         st.error(f"Error reading PDF file: {e}")
#         return ""

# def read_docx(file):
#     try:
#         doc = docx.Document(file)
#         full_text = []
#         for para in doc.paragraphs:
#             full_text.append(para.text)
#         return '\n'.join(full_text)
#     except Exception as e:
#         st.error(f"Error reading DOCX file: {e}")
#         return ""

# def analyze_and_respond(file_content):
#     response = get_gemini_response(file_content, prompt)
#     return response

# def get_greeting():
#     current_hour = datetime.now().hour
#     if current_hour < 12:
#         return "Good Morning", "#FF5733"  # Morning in orange
#     elif 12 <= current_hour < 18:
#         return "Good Afternoon", "#33C4FF"  # Afternoon in blue
#     elif 18 <= current_hour < 21:
#         return "Good Evening", "#8D33FF"  # Evening in purple
#     else:
#         return "Good Night", "#FF33A1"  # Night in pink

# def app():
#     st.markdown("""
#     <style>
#     .main-title {
#         font-family: 'Arial', sans-serif;
#         color: #2E4053;
#         text-align: center;
#         margin-bottom: 20px;
#     }
#     .greeting {
#         font-family: 'Arial', sans-serif;
#         text-align: center;
#         margin-top: 10px;
#         font-size: 24px;
#     }
#     .section-title {
#         font-family: 'Arial', sans-serif;
#         color: white;
#         margin-top: 20px;
#         margin-bottom: 10px;
#     }
#     </style>
#     """, unsafe_allow_html=True)

#     st.markdown('<h1 class="main-title">Welcome to Chatlop</h1>', unsafe_allow_html=True)

#     # Display greeting based on time of day
#     greeting, color = get_greeting()
#     st.markdown(f'<h3 class="greeting" style="color:{color};">{greeting}!</h3>', unsafe_allow_html=True)

#     option = st.sidebar.selectbox(
#         'Choose an option',
#         ['Chat with Bot', 'Upload Document']
#     )

#     if option == 'Chat with Bot':
#         st.markdown('<h2 class="section-title">Chat with the Bot 🤖</h2>', unsafe_allow_html=True)
#         if "messages" not in st.session_state:
#             st.session_state["messages"] = [{"role": "assistant", "content": "How can I help you?"}]
        
#         for msg in st.session_state.messages:
#             st.chat_message(msg["role"]).write(msg["content"])
        
#         if prompt := st.chat_input():
#             st.session_state.messages.append({"role": "user", "content": prompt})
#             st.chat_message("user").write(prompt)
#             response = get_gemini_response(prompt, prompt)
#             if response:
#                 msg = response
#                 st.session_state.messages.append({"role": "assistant", "content": msg})
#                 st.chat_message("assistant").write(msg)
#             else:
#                 st.error("No valid response could be generated.")

#     elif option == 'Upload Document':
#         st.markdown('<h2 class="section-title">Upload a File 📔</h2>', unsafe_allow_html=True)
#         st.markdown('<div class="upload-section">', unsafe_allow_html=True)
#         uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx'])
#         st.markdown('</div>', unsafe_allow_html=True)
        
#         if st.button("Analyze File", key='file_analyze', help="Analyze the uploaded file"):
#             if uploaded_file:
#                 file_content = read_uploaded_file(uploaded_file)
#                 if file_content:
#                     response = analyze_and_respond(file_content)
#                     if response:
#                         st.markdown('<div class="file-content">{}</div>'.format(response), unsafe_allow_html=True)
#                     else:
#                         st.error("No valid response could be generated from the file content.")
#             else:
#                 st.error("Please upload a file.")

#     elif option == 'History':
#         st.markdown('<h2 class="section-title">History 📜</h2>', unsafe_allow_html=True)
#         st.write("This section can be used to display the history of interactions or uploaded files.")

# app()


import streamlit as st
import google.generativeai as genai
from gtts import gTTS
from pydub import AudioSegment
from pydub.playback import play
import tempfile
import os
import fitz  # PyMuPDF for PDF files
import docx  # python-docx for DOCX files
from datetime import datetime

# Replace with your actual API key or use environment variables
GOOGLE_API_KEY = 'AIzaSyATZ0NDoTweYu3S3KQBlLBo0PljHvViK30'
genai.configure(api_key=GOOGLE_API_KEY)

st.set_page_config(page_title="Chatlop", page_icon="🤖", layout="wide")

# Prompts for different analysis tasks
ANALYSIS_PROMPTS = {
    "key_points": """
    Analyze the following document and extract the top 5-7 most important key points. 
    Provide a concise summary that captures the essential information in a clear, structured manner.
    """,
    
    "potential_questions": """
    Based on the document content, generate 5-7 potential questions that a reader might ask. 
    These questions should cover the most critical and intriguing aspects of the document.
    Provide both the questions and brief, informative answers.
    """,
    
    "summary": """
    Provide a comprehensive summary of the document. 
    The summary should capture the main ideas, key arguments, and overall context of the text.
    Aim for a summary that is both detailed and concise, highlighting the most significant information.
    """
}

def get_gemini_response(content, prompt):
    """Generate response using Gemini AI"""
    try:
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content([prompt, content])
        return response.text
    except Exception as e:
        st.error(f"Error getting response from GenAI: {e}")
        return None

def read_uploaded_file(file):
    """Read contents of uploaded PDF or DOCX file"""
    try:
        if file is not None:
            file_type = file.type
            if 'pdf' in file_type:
                return read_pdf(file)
            elif 'wordprocessingml' in file_type:
                return read_docx(file)
            else:
                st.error("Unsupported file type.")
                return None
        return None
    except Exception as e:
        st.error(f"Error reading uploaded file: {e}")
        return None

def read_pdf(file):
    """Read PDF file contents"""
    try:
        pdf_document = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text()
        return text
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return ""

def read_docx(file):
    """Read DOCX file contents"""
    try:
        doc = docx.Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return ""

def perform_comprehensive_analysis(file_content):
    """Perform comprehensive document analysis"""
    analysis_results = {}
    
    # Perform different types of analysis
    for analysis_type, prompt in ANALYSIS_PROMPTS.items():
        response = get_gemini_response(file_content, prompt)
        analysis_results[analysis_type] = response
    
    return analysis_results

def display_analysis_results(analysis_results):
    """Display analysis results in Streamlit"""
    # Key Points Section
    st.markdown("### 📌 Key Points")
    st.markdown(analysis_results.get('key_points', 'No key points found.'))
    
    # Potential Questions Section
    st.markdown("### ❓ Potential Questions and Answers")
    st.markdown(analysis_results.get('potential_questions', 'No questions generated.'))
    
    # Summary Section
    st.markdown("### 📖 Document Summary")
    st.markdown(analysis_results.get('summary', 'No summary available.'))

def get_greeting():
    """Get greeting based on time of day"""
    current_hour = datetime.now().hour
    if current_hour < 12:
        return "Good Morning", "#FF5733"  # Morning in orange
    elif 12 <= current_hour < 18:
        return "Good Afternoon", "#33C4FF"  # Afternoon in blue
    elif 18 <= current_hour < 21:
        return "Good Evening", "#8D33FF"  # Evening in purple
    else:
        return "Good Night", "#FF33A1"  # Night in pink

def app():
    # Custom CSS
    st.markdown("""
    <style>
    .main-title {
        font-family: 'Arial', sans-serif;
        color: #2E4053;
        text-align: center;
        margin-bottom: 20px;
    }
    .greeting {
        font-family: 'Arial', sans-serif;
        text-align: center;
        margin-top: 10px;
        font-size: 24px;
    }
    .section-title {
        font-family: 'Arial', sans-serif;
        color: white;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    .analysis-section {
        background-color: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
        margin-top: 20px;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<h1 class="main-title">Welcome to Chatlop</h1>', unsafe_allow_html=True)

    # Display greeting based on time of day
    greeting, color = get_greeting()
    st.markdown(f'<h3 class="greeting" style="color:{color};">{greeting}!</h3>', unsafe_allow_html=True)

    # Main application
    st.sidebar.title("Document Analysis")
    option = st.sidebar.selectbox(
        'Choose an option',
        ['Chat with Bot', 'Document Analysis']
    )

    if option == 'Chat with Bot':
        st.markdown('<h2 class="section-title">Chat with the Bot 🤖</h2>', unsafe_allow_html=True)
        if "messages" not in st.session_state:
            st.session_state["messages"] = [{"role": "assistant", "content": "How can I help you?"}]
        
        for msg in st.session_state.messages:
            st.chat_message(msg["role"]).write(msg["content"])
        
        if prompt := st.chat_input():
            st.session_state.messages.append({"role": "user", "content": prompt})
            st.chat_message("user").write(prompt)
            response = get_gemini_response(prompt, prompt)
            if response:
                msg = response
                st.session_state.messages.append({"role": "assistant", "content": msg})
                st.chat_message("assistant").write(msg)
            else:
                st.error("No valid response could be generated.")

    elif option == 'Document Analysis':
        st.markdown('<h2 class="section-title">Document Analysis 📔</h2>', unsafe_allow_html=True)
        
        # File Upload Section
        uploaded_file = st.file_uploader(
            "Choose a file", 
            type=['pdf', 'docx'], 
            help="Upload a PDF or Word document for comprehensive analysis"
        )
        
        # Analysis Options
        analysis_options = st.multiselect(
            "Select Analysis Types",
            ["Key Points", "Potential Questions", "Summary"],
            default=["Key Points", "Potential Questions", "Summary"]
        )
        
        # Analyze Button
        if st.button("Analyze Document", key='file_analyze', help="Perform comprehensive document analysis"):
            if uploaded_file:
                # Read file contents
                file_content = read_uploaded_file(uploaded_file)
                
                if file_content:
                    # Perform analysis
                    with st.spinner('Analyzing document...'):
                        analysis_results = perform_comprehensive_analysis(file_content)
                    
                    # Display results in a styled section
                    st.markdown('<div class="analysis-section">', unsafe_allow_html=True)
                    
                    # Only display selected analysis types
                    if "Key Points" in analysis_options:
                        st.markdown("### 📌 Key Points")
                        st.write(analysis_results.get('key_points', 'No key points found.'))
                    
                    if "Potential Questions" in analysis_options:
                        st.markdown("### ❓ Potential Questions and Answers")
                        st.write(analysis_results.get('potential_questions', 'No questions generated.'))
                    
                    if "Summary" in analysis_options:
                        st.markdown("### 📖 Document Summary")
                        st.write(analysis_results.get('summary', 'No summary available.'))
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.error("Unable to read the uploaded file.")
            else:
                st.error("Please upload a PDF or Word document.")

# Run the app
app()
